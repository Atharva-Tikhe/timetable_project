import sys

import docx
import pandas as pd
from PyQt5.QtWidgets import *
# import docx
# import os
# from actual_design import *
from responsive_test import *
import sqlite3
import datetime

year = ''
sem = ''
program = ''
curr_date = ''
push_count = 0
selectedDays = []
selectedDurations = ['']
table_data = []
preview_table_items = []
preview_table_header_items = []

conn = sqlite3.connect('courses_and_subjects_2.db')
cursor = conn.cursor()
conn.commit()

class TtInterface(QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.durationLabel1A.setVisible(False)
        self.ui.durationLabel2A.setVisible(False)
        self.ui.startTimeA.setVisible(False)
        self.ui.endTimeA.setVisible(False)
        self.ui.durationLabel1B.setVisible(False)
        self.ui.durationLabel2B.setVisible(False)
        self.ui.startTimeB.setVisible(False)
        self.ui.endTimeB.setVisible(False)
        self.ui.auto_fill_end_sem_time.setVisible(False)
        self.ui.examType.addItems(['Choose Exam','Mid Sem','End Sem'])
        self.ui.program_widget.addItems(['B.Tech', 'M.Tech'])
        self.ui.semComboBox.addItems(['Choose Sem','1','2','3','4','5','6','7','8','9','10'])
        self.ui.horizontalSlider.setMinimum(datetime.date.today().year - 10)
        self.ui.horizontalSlider.setMaximum(datetime.date.today().year + 10)
        self.ui.horizontalSlider.setSingleStep(1)
        self.ui.program_widget.itemClicked.connect(self.get_program)
        self.ui.horizontalSlider.valueChanged.connect(self.show_slider_value)
        self.ui.semComboBox.currentIndexChanged.connect(self.set_year_and_sem)
        self.ui.pushButton.clicked.connect(self.parse_db)
        self.ui.pushButton.clicked.connect(self.exam)
        self.ui.calendar.selectionChanged.connect(self.get_date)
        self.ui.pushToTable.clicked.connect(self.push_to_table)
        self.ui.clear_btn.clicked.connect(self.clear_everything)
        self.ui.save.clicked.connect(self.get_data)
        self.ui.course_code_selection.clicked.connect(self.auto_fill_code)


    def get_data(self):
        '''Stores data from the preview table into a global list'''
        global table_data
        each_row = []

        for i in range(self.ui.preview_table.rowCount()):
            for j in range(self.ui.preview_table.columnCount()):
                each_row.append(self.ui.preview_table.item(i,j).text())

        for i in range(0, len(each_row), 4):
            table_data.append(each_row[i:i + 4])
        print(table_data)

        preview_df = pd.DataFrame(i for i in table_data)
        preview_df.columns = ["Course Code", "Course Name", "Day and Date", "Time"]
        doc = docx.Document()
        doc.add_heading(f"Exam Schedule {sem} semester \n ({self.ui.batch_text_preview.text()} Pattern)")

        table = doc.add_table(preview_df.shape[0]+1, preview_df.shape[1])
        table.style("grid")
        for j in range(preview_df.shape[-1]):
            table.cell(0,j).text = preview_df.columns[j]
        for i in range(preview_df.shape[0]):
            for j in range(preview_df.shape[-1]):
                table.cell(i+1,j).text = str(preview_df.values[i,j])

        doc.save("timetable.docx")


    def clear_everything(self):
        '''Clears every filled widget'''
        self.ui.preview_table.clearContents()
        self.ui.subjectListWidget.clear()
        self.ui.program_widget.clearSelection()
        self.clear_one_thread()


    def clear_one_thread(self):
        '''Extension of clear_everything() since running on one thread'''
        self.ui.time_list.clear()


    def auto_fill_code(self):
        '''Fills the course code by looking up the subjects in the database'''
        global program
        code_dict = {}
        #TODO: remove the method, work for radiobutton #DONE
        if program == 'B.Tech':
            batch_shorthand = 'B'
        else:
            batch_shorthand = 'IM'
        final_code = str(self.ui.batch_text_preview.text()[2:]) + batch_shorthand
        print(final_code)

        for i in range(self.ui.preview_table.rowCount()):
            code_dict[self.ui.preview_table.item(i,1).text()] = ''
        print(code_dict)

        for i in code_dict:
            code_dict[i] = final_code + cursor.execute(f'''SELECT course_code FROM course_code
                                                WHERE subject_name = '{i}' ''').fetchall()[0][0]

        for i in range(self.ui.preview_table.rowCount()):
            self.ui.preview_table.setItem(i,0,QTableWidgetItem(code_dict[self.ui.preview_table.item(i, 1).text()]))



    def get_program(self):
        '''Gets program name and sets it to global var 'program' '''
        global program
        program = self.ui.program_widget.currentItem().text()

    def show_slider_value(self):
        '''Fills a text field with the slider value(year)'''
        slider_value = self.ui.horizontalSlider.value()
        batch_text = str(slider_value)
        self.ui.batch_text_preview.setText(batch_text)

    def init_table(self):
        '''Initializes preview table and sets row counts and heights'''
        self.ui.preview_table.setRowCount(6)
        for i in range(self.ui.preview_table.rowCount()):
            self.ui.preview_table.setRowHeight(i,50)

    def set_year_and_sem(self):
        '''sets global variables `year` and `sem` as per the selection'''
        #TODO: Keep only semesters (1-8,10) and remove year dropdown #DONE
        global sem
        sem = self.ui.semComboBox.itemText(self.ui.semComboBox.currentIndex())


    def parse_db(self):
        global conn
        global cursor
        global sem
        parse = []
        #TODO: Change the entire method to suit the new database format. #DONE
        try:
            query_result = cursor.execute(f'''SELECT subject_name FROM course_code WHERE semester = {sem}''')
            result_set = query_result.fetchall()
            for i in range(len(result_set)):
                parse.append(result_set[i][0])

            self.ui.subjectListWidget.addItems(parse)
            self.ui.subjectListWidget.setDragDropMode(QAbstractItemView.DragOnly)
            self.ui.preview_table.setDragDropMode(QAbstractItemView.DragDrop)
        except sqlite3.OperationalError:
            self.ui.statusBar.showMessage("Error while fetching sql results, try drop down again.", 500)


        # self.clear() #TODO Check for working

    @staticmethod
    def clear():
        '''clears the year and sem global variables'''
        global year
        global sem
        year = ''
        sem = ''

    def exam(self):
        '''selectively shows number of possible durations based on exam type(midsem/endsem)'''
        if self.ui.examType.currentIndexChanged:
            if self.ui.examType.currentIndex() == 1:
                self.ui.durationLabel1A.setVisible(True)
                self.ui.durationLabel2A.setVisible(True)
                self.ui.startTimeA.setVisible(True)
                self.ui.endTimeA.setVisible(True)
                self.ui.durationLabel1B.setVisible(True)
                self.ui.durationLabel2B.setVisible(True)
                self.ui.startTimeB.setVisible(True)
                self.ui.endTimeB.setVisible(True)
            if self.ui.examType.currentIndex() == 2:
                self.ui.durationLabel1A.setVisible(True)
                self.ui.durationLabel2A.setVisible(True)
                self.ui.startTimeA.setVisible(True)
                self.ui.endTimeA.setVisible(True)
                self.ui.durationLabel1B.setVisible(False)
                self.ui.durationLabel2B.setVisible(False)
                self.ui.startTimeB.setVisible(False)
                self.ui.endTimeB.setVisible(False)
            if self.ui.examType.currentIndex() == 0:
                self.ui.durationLabel1A.setVisible(False)
                self.ui.durationLabel2A.setVisible(False)
                self.ui.startTimeA.setVisible(False)
                self.ui.endTimeA.setVisible(False)
                self.ui.durationLabel1B.setVisible(False)
                self.ui.durationLabel2B.setVisible(False)
                self.ui.startTimeB.setVisible(False)
                self.ui.endTimeB.setVisible(False)


    def get_date(self):
        '''Gets selected dates from the calendar and appends to a global list'''
        global curr_date
        y,m,d = self.ui.calendar.selectedDate().getDate()
        dayName = self.ui.calendar.selectedDate().longDayName(self.ui.calendar.selectedDate().dayOfWeek())
        curr_date = f'{dayName} \n {d}/{m}/{y}'
        selectedDays.append(curr_date)
        if self.ui.examType.currentIndex() == 1:
            selectedDays.append(curr_date)


    def push_to_table(self):
        global curr_date
        global selectedDays
        global selectedDurations
        self.init_table()
        if self.ui.examType.currentIndex() == 1:
            duration = f'{self.ui.startTimeA.time().toString("h:mm ap")} - \n {self.ui.endTimeA.time().toString("h:mm ap")} '
            selectedDurations.append(duration)
            durationB = f'{self.ui.startTimeB.time().toString("h:mm ap")} - \n {self.ui.endTimeB.time().toString("h:mm ap")} '
            selectedDurations.append(durationB)
            for i in range(len(selectedDays)):
                self.ui.preview_table.setItem(i, 2, QTableWidgetItem(selectedDays[i]))

        elif self.ui.examType.currentIndex() == 2:
            duration = f'{self.ui.startTimeA.time().toString("h:mm ap")} - \n {self.ui.endTimeA.time().toString("h:mm ap")} '
            selectedDurations.append(duration)
            for i in range(len(selectedDays)):
                    self.ui.preview_table.setItem(i, 2, QTableWidgetItem(selectedDays[i]))
            selectedDurations.remove('')
            for j in range(len(selectedDurations)):
                self.ui.preview_table.setItem(j,3, QTableWidgetItem(selectedDurations[j]))
            self.ui.auto_fill_end_sem_time.setEnabled(True)
            self.ui.auto_fill_end_sem_time.setVisible(True)
            self.ui.auto_fill_end_sem_time.clicked.connect(self.auto_fill_time)

        self.ui.time_list.setDragDropMode(QAbstractItemView.DragOnly)
        self.ui.time_list.addItems(selectedDurations)

    def auto_fill_time(self):
        global selectedDurations
        for i in range(self.ui.preview_table.rowCount()):
            self.ui.preview_table.setItem(i,3, QTableWidgetItem(selectedDurations[0]))




if __name__ == "__main__":
    app = QApplication(sys.argv)
    w = TtInterface()
    w.show()
    sys.exit(app.exec_())

